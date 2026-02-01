import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
from datetime import datetime
import io
import zipfile
from typing import Dict, List, Tuple
import re

def parse_date(date_str: str) -> datetime:
    """Parse date string in format 'MMM DD, YYYY' to datetime object"""
    try:
        return datetime.strptime(date_str, "%b %d, %Y")
    except ValueError:
        # Fallback for different date formats
        try:
            return datetime.strptime(date_str, "%B %d, %Y")
        except ValueError:
            # If all else fails, return a default date
            return datetime.min

def replace_mp_name_in_letter(letter_content: str, salutation: str, last_name: str) -> str:
    """Replace [MP Name] placeholder with salutation and last name"""
    mp_greeting = f"{salutation} {last_name}"
    return letter_content.replace("[MP Name]", mp_greeting)

def create_docx_for_electorate(electorate: str, mp_info: Dict, letters: List[Dict]) -> io.BytesIO:
    """Create a DOCX document for a specific electorate with all letters"""
    doc = Document()
    
    # Set default font to Roboto for the entire document
    style = doc.styles['Normal']
    style.font.name = 'Roboto'
    
    # Sort letters by date (ascending)
    sorted_letters = sorted(letters, key=lambda x: parse_date(x['Submission Date']))
    
    for i, letter in enumerate(sorted_letters):
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.add_run(f"Date: {letter['Submission Date']}").bold = True
        
        # Add empty line
        doc.add_paragraph()
        
        # Process letter content - replace MP name
        letter_content = replace_mp_name_in_letter(
            letter['Your letter'], 
            mp_info['Salutation'], 
            mp_info['Last name']
        )
        
        # Add letter content
        doc.add_paragraph(letter_content)
        
        # Add empty line
        doc.add_paragraph()
        
        # Add postcode and state
        location_paragraph = doc.add_paragraph()
        location_paragraph.add_run(f"{letter['POSTCODE']}, {letter['STATE']}").bold = True
        
        # Add page break if not the last letter
        if i < len(sorted_letters) - 1:
            doc.add_page_break()
    
    # Save to BytesIO
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer

def process_files(mps_df: pd.DataFrame, letters_df: pd.DataFrame) -> Dict[str, io.BytesIO]:
    """Process the uploaded files and generate DOCX files for each electorate"""
    results = {}
    
    # Create a dictionary for MP lookup - handle multiple MPs per electorate
    mp_dict = {}
    for _, mp in mps_df.iterrows():
        electorate = mp['State/Electorate']
        full_name = f"{mp['First name']} {mp['Last name']}"
        mp_key = f"{electorate}_{full_name}"  # Unique key for each MP
        mp_dict[mp_key] = {
            'Electorate': electorate,
            'Salutation': mp['Salutation'],
            'First name': mp['First name'],
            'Last name': mp['Last name'],
            'Fullname': full_name  # Keep for filename compatibility
        }
    
    # Group letters by electorate and state (only for those that have MPs)
    electorate_letters = {}
    for _, letter in letters_df.iterrows():
        # Check ELECTORATE column
        electorate_raw = str(letter['ELECTORATE'])
        
        # Handle comma-separated and newline-separated electorates 
        # (e.g., "Bruce, Hotham" or "Blaxland\nMcMahon")
        # First split by newlines, then by commas
        temp_electorates = []
        for part in electorate_raw.split('\n'):
            temp_electorates.extend([e.strip() for e in part.split(',')])
        
        electorates = [e.strip() for e in temp_electorates if e.strip()]
        
        # Also check STATE column
        state = str(letter['STATE']).strip()
        if state:
            electorates.append(state)
        
        for electorate in electorates:
            # Find all MPs for this electorate/state
            matching_mps = [key for key in mp_dict.keys() if mp_dict[key]['Electorate'] == electorate]
            
            # Add letter to each MP's document for this electorate
            for mp_key in matching_mps:
                if mp_key not in electorate_letters:
                    electorate_letters[mp_key] = []
                
                electorate_letters[mp_key].append({
                    'Submission Date': letter['Submission Date'],
                    'Your letter': letter['Your letter'],
                    'POSTCODE': letter['POSTCODE'],
                    'STATE': letter['STATE']
                })
    
    # Generate DOCX for each MP that has letters
    for mp_key in mp_dict.keys():
        if mp_key in electorate_letters:
            letters = electorate_letters[mp_key]
            mp_info = mp_dict[mp_key]
            
            doc_buffer = create_docx_for_electorate(mp_info['Electorate'], mp_info, letters)
            results[mp_key] = doc_buffer
    
    return results

def create_zip_file(docx_files: Dict[str, io.BytesIO], mp_dict: Dict) -> io.BytesIO:
    """Create a ZIP file containing all DOCX files"""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for mp_key, doc_buffer in docx_files.items():
            # Get MP info for filename
            mp_info = mp_dict.get(mp_key, {})
            electorate = mp_info.get('Electorate', 'Unknown')
            fullname = mp_info.get('Fullname', 'Unknown')
            
            # Create filename: "Electorate, Fullname.docx"
            clean_electorate = re.sub(r'[<>:"/\\|?*]', '', electorate)
            clean_fullname = re.sub(r'[<>:"/\\|?*]', '', fullname)
            filename = f"{clean_electorate}, {clean_fullname}.docx"
            zip_file.writestr(filename, doc_buffer.getvalue())
    
    zip_buffer.seek(0)
    return zip_buffer

def main():
    st.title("MP Letters to DOCX Converter")
    st.write("Upload MP and Letters CSV files to generate formatted DOCX documents for each electorate.")
    
    # File upload section
    st.header("Upload Files")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("MPs CSV File")
        mps_file = st.file_uploader(
            "Upload MPs CSV", 
            type=['csv'], 
            key="mps_file",
            help="CSV file containing MP information with columns: Salutation, Fullname, State/Electorate"
        )
    
    with col2:
        st.subheader("Letters CSV File")
        letters_file = st.file_uploader(
            "Upload Letters CSV", 
            type=['csv'], 
            key="letters_file",
            help="CSV file containing letters with columns: ELECTORATE, Submission Date, Your letter, POSTCODE, STATE"
        )
    
    if mps_file is not None and letters_file is not None:
        try:
            # Load the CSV files
            mps_df = pd.read_csv(mps_file)
            letters_df = pd.read_csv(letters_file)
            
            # Validate required columns
            required_mp_columns = ['Salutation', 'First name', 'Last name', 'State/Electorate']
            required_letter_columns = ['ELECTORATE', 'Submission Date', 'Your letter', 'POSTCODE', 'STATE']
            
            missing_mp_cols = [col for col in required_mp_columns if col not in mps_df.columns]
            missing_letter_cols = [col for col in required_letter_columns if col not in letters_df.columns]
            
            if missing_mp_cols:
                st.error(f"MPs CSV is missing required columns: {', '.join(missing_mp_cols)}")
                return
            
            if missing_letter_cols:
                st.error(f"Letters CSV is missing required columns: {', '.join(missing_letter_cols)}")
                return
            
            # Display file information
            st.success("Files uploaded successfully!")
            
            col1, col2 = st.columns(2)
            with col1:
                st.info(f"**MPs loaded:** {len(mps_df)} records")
                st.write("Sample MP data:")
                st.dataframe(mps_df.head(3))
            
            with col2:
                st.info(f"**Letters loaded:** {len(letters_df)} records")
                st.write("Sample Letter data:")
                st.dataframe(letters_df[['ELECTORATE', 'Submission Date', 'POSTCODE', 'STATE']].head(3))
            
            # Process files and generate documents
            if st.button("Generate DOCX Files", type="primary"):
                with st.spinner("Processing files and generating documents..."):
                    try:
                        docx_files = process_files(mps_df, letters_df)
                        
                        if not docx_files:
                            st.warning("No matching electorates found between MPs and Letters files.")
                            return
                        
                        # Create MP dict for filename generation (same structure as process_files)
                        mp_dict = {}
                        for _, mp in mps_df.iterrows():
                            electorate = mp['State/Electorate']
                            full_name = f"{mp['First name']} {mp['Last name']}"
                            mp_key = f"{electorate}_{full_name}"
                            mp_dict[mp_key] = {
                                'Electorate': electorate,
                                'Salutation': mp['Salutation'],
                                'First name': mp['First name'],
                                'Last name': mp['Last name'],
                                'Fullname': full_name
                            }
                        
                        st.success(f"Generated {len(docx_files)} DOCX files!")
                        
                        # Show summary
                        st.header("Generation Summary")
                        summary_data = []
                        
                        # Count letters by electorate and state (matching the processing logic)
                        letter_counts = {}
                        for _, letter in letters_df.iterrows():
                            # Check ELECTORATE column
                            electorate_raw = str(letter['ELECTORATE'])
                            
                            # Handle comma-separated and newline-separated electorates 
                            temp_electorates = []
                            for part in electorate_raw.split('\n'):
                                temp_electorates.extend([e.strip() for e in part.split(',')])
                            
                            electorates = [e.strip() for e in temp_electorates if e.strip()]
                            
                            # Also check STATE column
                            state = str(letter['STATE']).strip()
                            if state:
                                electorates.append(state)
                            
                            for electorate in electorates:
                                if electorate in letter_counts:
                                    letter_counts[electorate] += 1
                                else:
                                    letter_counts[electorate] = 1
                        
                        for mp_key in docx_files.keys():
                            mp_info = mp_dict.get(mp_key, {})
                            electorate = mp_info.get('Electorate', 'Unknown')
                            fullname = mp_info.get('Fullname', 'Unknown')
                            letter_count = letter_counts.get(electorate, 0)
                            summary_data.append({
                                'Electorate': electorate,
                                'MP': fullname,
                                'Letters': letter_count
                            })
                        
                        summary_df = pd.DataFrame(summary_data)
                        st.dataframe(summary_df, width='stretch')
                        
                        # Create download options
                        st.header("Download Options")
                        
                        # Individual file downloads
                        with st.expander("Download Individual Files"):
                            for mp_key, doc_buffer in docx_files.items():
                                # Get MP info for filename
                                mp_info = mp_dict.get(mp_key, {})
                                electorate = mp_info.get('Electorate', 'Unknown')
                                fullname = mp_info.get('Fullname', 'Unknown')
                                
                                # Create filename: "Electorate, Fullname.docx"
                                clean_electorate = re.sub(r'[<>:"/\\|?*]', '', electorate)
                                clean_fullname = re.sub(r'[<>:"/\\|?*]', '', fullname)
                                filename = f"{clean_electorate}, {clean_fullname}.docx"
                                
                                st.download_button(
                                    label=f"📄 Download {electorate} - {fullname}",
                                    data=doc_buffer.getvalue(),
                                    file_name=filename,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                        
                        # Batch download as ZIP
                        st.subheader("Batch Download")
                        zip_buffer = create_zip_file(docx_files, mp_dict)
                        
                        st.download_button(
                            label="📦 Download All Files as ZIP",
                            data=zip_buffer.getvalue(),
                            file_name="all_electorate_letters.zip",
                            mime="application/zip"
                        )
                        
                    except Exception as e:
                        st.error(f"Error processing files: {str(e)}")
                        st.write("Please check your CSV file formats and try again.")
        
        except Exception as e:
            st.error(f"Error loading CSV files: {str(e)}")
            st.write("Please ensure your files are valid CSV format.")
    
    else:
        st.info("Please upload both CSV files to begin processing.")
        
        # Show expected file formats
        with st.expander("Expected File Formats"):
            st.subheader("MPs CSV Format")
            st.write("Required columns:")
            st.code("""
Salutation,First name,Last name,State/Electorate
Mr,David,Smith,Bean
Senator,Jennifer,McAllister,NSW
            """)
            
            st.subheader("Letters CSV Format")
            st.write("Required columns:")
            st.code("""
ELECTORATE,Submission Date,Your letter,POSTCODE,STATE
Aston,"Jun 28, 2025","Dear [MP Name], ...",4212,QLD
Bennelong,"Jul 10, 2025","Dear [MP Name], ...",2154,NSW
            """)

if __name__ == "__main__":
    main()
